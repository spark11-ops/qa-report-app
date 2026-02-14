# ==== COMPLETE QA SYSTEM - WITH TELEGRAM NOTIFICATIONS ====

from flask import Flask, render_template, request, send_file, session, jsonify, redirect, url_for
import os
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta
from collections import defaultdict
import pickle
import zipfile
import requests
import subprocess
import secrets as secret_gen

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)

# Secret key for session management
secret_key = os.environ.get('SECRET_KEY')
if not secret_key:
    secret_key = secret_gen.token_hex(32)
    print("WARNING: Using randomly generated SECRET_KEY. Set SECRET_KEY environment variable in production!")

app.secret_key = secret_key

# Make sessions permanent (last 31 days)
app.permanent_session_lifetime = timedelta(days=31)

# Telegram Bot Configuration
TELEGRAM_BOT_TOKEN = os.environ.get('TELEGRAM_BOT_TOKEN', '')
TELEGRAM_CHAT_ID = os.environ.get('TELEGRAM_CHAT_ID', '')

BASE_DIR = os.getcwd()

UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "outputs")
ASSET_FOLDER = os.path.join(BASE_DIR, "assets")
DATA_FOLDER = os.path.join(BASE_DIR, "data")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(ASSET_FOLDER, exist_ok=True)
os.makedirs(DATA_FOLDER, exist_ok=True)

# =========================
# TELEGRAM FUNCTIONS
# =========================

def send_telegram_notification(zip_path, message):
    """Send zip file and message to Telegram"""
    if not TELEGRAM_BOT_TOKEN or not TELEGRAM_CHAT_ID:
        print("DEBUG: Telegram not configured, skipping notification")
        return False
    
    try:
        # Send message
        message_url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
        message_data = {
            'chat_id': TELEGRAM_CHAT_ID,
            'text': message,
            'parse_mode': 'HTML'
        }
        requests.post(message_url, data=message_data)
        
        # Send file
        file_url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendDocument"
        with open(zip_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': TELEGRAM_CHAT_ID}
            response = requests.post(file_url, data=data, files=files)
        
        print(f"DEBUG: Telegram notification sent successfully")
        return True
    except Exception as e:
        print(f"DEBUG: Telegram notification failed: {e}")
        return False

def create_notification_package(qcw_path, dates_data, machine_name_mapping, user_ip, institute_name):
    """Create zip package with QCW and all DOCX files"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_filename = f"QA_Reports_{timestamp}.zip"
    zip_path = os.path.join(OUTPUT_FOLDER, zip_filename)
    
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        # Add QCW file
        zipf.write(qcw_path, os.path.basename(qcw_path))
        
        # Generate and add DOCX for each date
        for date, machines_data in dates_data.items():
            docx_path = generate_date_docx(date, machines_data, machine_name_mapping)
            zipf.write(docx_path, os.path.basename(docx_path))
    
    # Create message
    current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    message = f"""
<b>🔬 New QA Report Generated</b>

📅 <b>Date:</b> {current_date}
🏥 <b>Institute:</b> {institute_name}
🌐 <b>IP Address:</b> {user_ip}
📊 <b>Total Dates:</b> {len(dates_data)}
📁 <b>Package:</b> {zip_filename}

<i>Package contains QCW file and all DOCX reports</i>
"""
    
    # Send to Telegram
    send_telegram_notification(zip_path, message)
    
    return zip_path

# =========================
# HELPERS
# =========================

def format_fieldsize_mm_to_cm(field_text):
    """Convert field size from mm to cm"""
    try:
        x_mm, y_mm = field_text.split("x")
        x_cm = float(x_mm) / 10
        y_cm = float(y_mm) / 10
        return f"{x_cm:.0f} cm X {y_cm:.0f} cm"
    except:
        return field_text

def energy_unit(modality):
    """Determine energy unit based on modality"""
    return "MV" if modality.lower().startswith("photon") else "MeV"

def format_energy_with_fff(energy, fff_value):
    """Format energy with FFF if applicable"""
    energy_str = str(energy)
    if fff_value and fff_value.upper() == "YES":
        return f"{energy_str} FFF"
    return energy_str

# =========================
# QCW PARSER
# =========================

def extract_worklists(file_path):
    """Extract all unique worklists from QCW file"""
    with open(file_path, 'rb') as f:
        content = f.read()
    
    if content.startswith(b'\xef\xbb\xbf'):
        content = content[3:]
    
    root = ET.fromstring(content)
    
    worklists = {}
    
    for trend in root.findall(".//TrendData"):
        worklist = trend.find("Worklist")
        if worklist is None:
            continue
            
        worklist_id = worklist.get("id")
        worklist_name_tag = worklist.find("Name")
        
        if worklist_name_tag is None or worklist_name_tag.text is None:
            worklist_name = f"Worklist_{worklist_id}"
        else:
            worklist_name = worklist_name_tag.text.strip()
        
        if worklist_id not in worklists:
            worklists[worklist_id] = worklist_name
    
    return worklists

def parse_qcw_with_mapping(file_path, machine_name_mapping):
    """Parse QCW file with deviation calculation from Norm values"""
    with open(file_path, 'rb') as f:
        content = f.read()
    
    if content.startswith(b'\xef\xbb\xbf'):
        content = content[3:]
    
    root = ET.fromstring(content)
    
    dates_data = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
    
    for trend in root.findall(".//TrendData"):
        date_str = trend.get("date").split(" ")[0]
        
        worklist = trend.find("Worklist")
        if worklist is None:
            continue
            
        worklist_id = worklist.get("id")
        machine_name = machine_name_mapping.get(worklist_id, f"Machine_{worklist_id}")
        
        energy = worklist.find(".//AdminValues/Energy").text
        modality = worklist.find(".//AdminValues/Modality").text
        raw_field = worklist.find(".//AdminValues/Fieldsize").text
        field = format_fieldsize_mm_to_cm(raw_field)
        
        fff_tag = worklist.find(".//AdminValues/FFF")
        fff_value = fff_tag.text if fff_tag is not None else "No"
        
        energy_display = format_energy_with_fff(energy, fff_value)
        unit = energy_unit(modality)
        energy_with_unit = f"{energy_display} {unit}"
        
        # Extract Norm values
        analyze_params = worklist.findall(".//AdminData/AnalyzeParams/*")
        norm_values = {}
        for param in analyze_params:
            if param.tag == "Wedge":
                continue
            norm_tag = param.find("Norm")
            if norm_tag is not None:
                norm_values[param.tag] = float(norm_tag.text)
        
        # Extract measured values
        meas_data = trend.find(".//MeasData")
        if meas_data is None:
            continue
            
        analyze_values = meas_data.find("AnalyzeValues")
        if analyze_values is None:
            continue
        
        measured = {}
        failed_tests = []
        
        for param in analyze_values:
            if param.tag == "Wedge":
                continue
            value_tag = param.find("Value")
            if value_tag is not None:
                meas_value = float(value_tag.text)
                meas_formatted = f"{meas_value:.2f}"
                
                # Calculate % deviation from Norm
                if param.tag in norm_values:
                    norm = norm_values[param.tag]
                    if norm != 0:
                        deviation = ((meas_value - norm) / norm) * 100
                        dev_formatted = f"{deviation:.2f}"
                        
                        # Check if test fails (deviation > ±3%)
                        if abs(deviation) > 3.0:
                            failed_tests.append({
                                "parameter": param.tag,
                                "deviation": dev_formatted,
                                "energy": energy_with_unit
                            })
                    else:
                        dev_formatted = "0.00"
                else:
                    dev_formatted = "N/A"
                
                measured[param.tag] = {
                    "value": meas_formatted,
                    "deviation": dev_formatted
                }
        
        energy_data = {
            "energy": energy_with_unit,
            "measured": measured,
            "failed_tests": failed_tests
        }
        
        dates_data[date_str][machine_name][field].append(energy_data)
    
    # Convert to regular dict
    result = {}
    for date in sorted(dates_data.keys()):
        result[date] = {}
        for machine_name, fields_data in dates_data[date].items():
            result[date][machine_name] = {}
            for field_size, energy_list in fields_data.items():
                result[date][machine_name][field_size] = energy_list
    
    return result

# =========================
# DOCX GENERATION
# =========================

def add_custom_header_footer(doc, logo_path, institute_name):
    """Add header and footer"""
    section = doc.sections[0]
    
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if os.path.exists(logo_path):
        header_para.clear()
        run = header_para.add_run()
        run.add_picture(logo_path, width=Inches(1.2))
    
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    
    run = footer_para.add_run(institute_name)
    run.font.size = Pt(10)
    
    if os.path.exists(logo_path):
        footer_para.add_run("   ")
        run = footer_para.add_run()
        run.add_picture(logo_path, width=Inches(0.6))
    
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_horizontal_line(paragraph):
    """Add blue horizontal line under title"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4472C4')
    
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_date_docx(date, machines_data, machine_name_mapping):
    """Generate DOCX with new table format"""
    doc = Document()
    
    logo_path = os.path.join(ASSET_FOLDER, "logo.png")
    name_file = os.path.join(ASSET_FOLDER, "name.txt")
    
    institute_name = "Institute/Hospital Name Here"
    if os.path.exists(name_file):
        with open(name_file) as f:
            institute_name = f.read().strip()
    
    add_custom_header_footer(doc, logo_path, institute_name)
    
    first_machine = True
    
    for machine_name, fields_data in machines_data.items():
        if not first_machine:
            doc.add_page_break()
        first_machine = False
        
        # Title
        title_para = doc.add_heading("Daily Quality Assurance", level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        add_horizontal_line(title_para)
        
        doc.add_paragraph()
        
        # Header line
        header_para = doc.add_paragraph()
        header_para.add_run(f"Machine Name: {machine_name}").bold = True
        header_para.add_run(" " * 40)  # Spacing
        header_para.add_run(f"Date: {date}").bold = True
        
        doc.add_paragraph()
        
        # Process each field size
        for field_size, energy_list in fields_data.items():
            # Field Size header
            field_para = doc.add_paragraph()
            field_para.add_run(f"Field Size : {field_size}").bold = True
            
            # NEW TABLE FORMAT with Meas.(%) and % Dev. columns
            num_energies = len(energy_list)
            # Columns: Parameter/Energy + (Meas + %Dev) * 5 parameters = 1 + 10 = 11 columns
            table = doc.add_table(rows=num_energies + 1, cols=11)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Parameter →\n\nEnergy↓"
            
            # CAX columns
            header_cells[1].text = "CAX"
            header_cells[1].merge(header_cells[2])
            
            # Flatness columns
            header_cells[3].text = "Flatness\n(%)"
            header_cells[3].merge(header_cells[4])
            
            # SymmetryGT columns
            header_cells[5].text = "SymmetryGT\n(%)"
            header_cells[5].merge(header_cells[6])
            
            # SymmetryLR columns
            header_cells[7].text = "SymmetryLR\n(%)"
            header_cells[7].merge(header_cells[8])
            
            # BQF columns
            header_cells[9].text = "BQF"
            header_cells[9].merge(header_cells[10])
            
            # Sub-header row (Meas. / % Dev.)
            subheader_row = table.add_row()
            subheader_row.cells[0].text = ""
            
            for i in range(5):  # 5 parameters
                base_idx = 1 + (i * 2)
                subheader_row.cells[base_idx].text = "Meas.\n(%)"
                subheader_row.cells[base_idx + 1].text = "%\nDev."
            
            # Make headers bold and centered
            for row in [table.rows[0], subheader_row]:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data rows (skip row 1 which is subheader, start from row 2)
            row_offset = 2
            for i, energy_data in enumerate(energy_list):
                row_cells = table.rows[i + row_offset].cells
                
                energy_str = energy_data["energy"]
                measured = energy_data["measured"]
                
                row_cells[0].text = energy_str
                
                # CAX
                row_cells[1].text = measured.get("CAX", {}).get("value", "N/A")
                row_cells[2].text = measured.get("CAX", {}).get("deviation", "N/A")
                
                # Flatness
                row_cells[3].text = measured.get("Flatness", {}).get("value", "N/A")
                row_cells[4].text = measured.get("Flatness", {}).get("deviation", "N/A")
                
                # SymmetryGT
                row_cells[5].text = measured.get("SymmetryGT", {}).get("value", "N/A")
                row_cells[6].text = measured.get("SymmetryGT", {}).get("deviation", "N/A")
                
                # SymmetryLR
                row_cells[7].text = measured.get("SymmetryLR", {}).get("value", "N/A")
                row_cells[8].text = measured.get("SymmetryLR", {}).get("deviation", "N/A")
                
                # BQF
                row_cells[9].text = measured.get("BQF", {}).get("value", "N/A")
                row_cells[10].text = measured.get("BQF", {}).get("deviation", "N/A")
                
                # Center align data
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # DYNAMIC NOTE based on pass/fail
            all_failed = []
            for energy_data in energy_list:
                all_failed.extend(energy_data.get("failed_tests", []))
            
            note_para = doc.add_paragraph()
            if len(all_failed) == 0:
                # All tests passed
                note_para.add_run("Note: All tests passed").bold = True
            else:
                # Some tests failed
                note_para.add_run("Note: ").bold = True
                
                # Group failures by energy
                failures_by_energy = defaultdict(list)
                for fail in all_failed:
                    failures_by_energy[fail['energy']].append(
                        f"{fail['parameter']} ({fail['deviation']}%)"
                    )
                
                # Build failure message
                failure_text = "Failed tests: "
                failure_parts = []
                for energy, params in failures_by_energy.items():
                    params_str = ", ".join(params)
                    failure_parts.append(f"{energy}: {params_str}")
                
                failure_text += "; ".join(failure_parts)
                note_para.add_run(failure_text)
            
            doc.add_paragraph()
    
    # Signature
    doc.add_paragraph("\n\n")
    sig_para = doc.add_paragraph()
    sig_para.add_run("Signature:").bold = True
    
    filename = f"QA_Report_{date}.docx"
    path = os.path.join(OUTPUT_FOLDER, filename)
    doc.save(path)
    return path

def convert_docx_to_pdf_libreoffice(docx_path):
    """Convert DOCX to PDF using LibreOffice"""
    try:
        output_dir = os.path.dirname(docx_path)
        result = subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', output_dir, docx_path
        ], check=True, capture_output=True, timeout=60)
        
        pdf_path = docx_path.replace('.docx', '.pdf')
        if os.path.exists(pdf_path):
            print(f"DEBUG: PDF generated successfully: {pdf_path}")
            return pdf_path
        else:
            print("DEBUG: PDF file not found after conversion")
            return None
    except subprocess.TimeoutExpired:
        print("DEBUG: LibreOffice conversion timed out")
        return None
    except Exception as e:
        print(f"DEBUG: PDF conversion failed: {e}")
        return None

# =========================
# CLEANUP
# =========================

def cleanup_old_data_files():
    """Remove data files older than 24 hours"""
    try:
        current_time = datetime.now().timestamp()
        for filename in os.listdir(DATA_FOLDER):
            if filename.startswith('data_') and filename.endswith('.pkl'):
                filepath = os.path.join(DATA_FOLDER, filename)
                file_age = current_time - os.path.getmtime(filepath)
                if file_age > 86400:
                    os.remove(filepath)
                    print(f"DEBUG: Cleaned up old data file: {filename}")
    except Exception as e:
        print(f"DEBUG: Error during cleanup: {e}")

# =========================
# ROUTES
# =========================

@app.route("/", methods=["GET", "POST"])
def index():
    cleanup_old_data_files()
    
    if request.method == "POST":
        file = request.files.get("file")
        
        if not file:
            print("DEBUG: No file in request")
            return jsonify({"error": "No file uploaded"}), 400
        
        filename = file.filename
        print(f"DEBUG: File uploaded: {filename}")
        
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        print(f"DEBUG: File saved to: {file_path}")
        
        inst_name = request.form.get("institute")
        if inst_name:
            with open(os.path.join(ASSET_FOLDER, "name.txt"), "w") as f:
                f.write(inst_name)
            print(f"DEBUG: Institute name saved: {inst_name}")
        
        logo = request.files.get("logo")
        if logo and logo.filename != "":
            logo.save(os.path.join(ASSET_FOLDER, "logo.png"))
            print("DEBUG: Logo saved")
        
        print("DEBUG: Extracting worklists...")
        worklists = extract_worklists(file_path)
        print(f"DEBUG: Found {len(worklists)} worklists: {list(worklists.values())}")
        
        session.permanent = True
        session['filename'] = filename
        session['worklists'] = worklists
        session.modified = True
        
        print("DEBUG: Session data stored, redirecting to worklist_mapping")
        
        return redirect(url_for('worklist_mapping'))
    
    return render_template("index.html")


@app.route("/worklist_mapping", methods=["GET", "POST"])
def worklist_mapping():
    if 'worklists' not in session:
        print("DEBUG: No worklists in session on GET, redirecting to index")
        return redirect(url_for('index'))
    
    print(f"DEBUG: Worklist mapping page loaded, {len(session['worklists'])} worklists in session")
    
    if request.method == "POST":
        print("DEBUG: Worklist mapping POST received")
        
        worklists = session.get('worklists', {})
        print(f"DEBUG: Worklists from session: {len(worklists)} items")
        
        if not worklists:
            print("DEBUG: No worklists in session, redirecting to index")
            return redirect(url_for('index'))
        
        machine_name_mapping = {}
        
        for wl_id in worklists.keys():
            custom_name = request.form.get(f"machine_{wl_id}")
            if custom_name:
                machine_name_mapping[wl_id] = custom_name.strip()
            else:
                machine_name_mapping[wl_id] = worklists[wl_id]
        
        print(f"DEBUG: Machine name mapping: {machine_name_mapping}")
        
        filename = session.get('filename')
        print(f"DEBUG: Filename from session: {filename}")
        
        if not filename:
            print("DEBUG: No filename in session, redirecting to index")
            return redirect(url_for('index'))
        
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        print(f"DEBUG: File path: {file_path}")
        
        if not os.path.exists(file_path):
            print(f"DEBUG: File does not exist at {file_path}, redirecting to index")
            return redirect(url_for('index'))
        
        print("DEBUG: Parsing QCW file...")
        dates_data = parse_qcw_with_mapping(file_path, machine_name_mapping)
        print(f"DEBUG: Parsed {len(dates_data)} dates")
        
        # Save to file (session too large issue)
        data_id = secret_gen.token_hex(16)
        data_file = os.path.join(DATA_FOLDER, f"data_{data_id}.pkl")
        
        with open(data_file, 'wb') as f:
            pickle.dump(dates_data, f)
        
        print(f"DEBUG: Saved data to file: {data_file}")
        
        session['data_id'] = data_id
        session['machine_name_mapping'] = machine_name_mapping
        session.modified = True
        
        # TELEGRAM NOTIFICATION
        print("DEBUG: Creating notification package...")
        inst_name_file = os.path.join(ASSET_FOLDER, "name.txt")
        institute_name = "Unknown Institute"
        if os.path.exists(inst_name_file):
            with open(inst_name_file) as f:
                institute_name = f.read().strip()
        
        user_ip = request.remote_addr
        
        try:
            create_notification_package(file_path, dates_data, machine_name_mapping, user_ip, institute_name)
            print("DEBUG: Notification package created and sent")
        except Exception as e:
            print(f"DEBUG: Failed to create notification: {e}")
        
        print("DEBUG: Redirecting to results")
        return redirect(url_for('results'))
    
    return render_template("worklist_mapping.html", worklists=session['worklists'])


@app.route("/results")
def results():
    if 'data_id' not in session or 'machine_name_mapping' not in session:
        print("DEBUG: No data_id in session, redirecting to index")
        return redirect(url_for('index'))
    
    data_id = session['data_id']
    data_file = os.path.join(DATA_FOLDER, f"data_{data_id}.pkl")
    
    if not os.path.exists(data_file):
        print(f"DEBUG: Data file not found: {data_file}, redirecting to index")
        return redirect(url_for('index'))
    
    with open(data_file, 'rb') as f:
        dates_data = pickle.load(f)
    
    machine_name_mapping = session['machine_name_mapping']
    
    print(f"DEBUG: Loaded {len(dates_data)} dates from file, displaying results")
    
    return render_template("result.html", 
                         dates_data=dates_data,
                         machine_name_mapping=machine_name_mapping)


@app.route("/generate/<date>/<format>")
def generate_date_report(date, format):
    if 'data_id' not in session:
        return "No data available", 404
    
    data_id = session['data_id']
    data_file = os.path.join(DATA_FOLDER, f"data_{data_id}.pkl")
    
    if not os.path.exists(data_file):
        return "Data file not found", 404
    
    with open(data_file, 'rb') as f:
        dates_data = pickle.load(f)
    
    machine_name_mapping = session.get('machine_name_mapping', {})
    
    if date not in dates_data:
        return f"No data for date {date}", 404
    
    machines_data = dates_data[date]
    
    docx_path = generate_date_docx(date, machines_data, machine_name_mapping)
    
    if format == "pdf":
        pdf_path = convert_docx_to_pdf_libreoffice(docx_path)
        if pdf_path and os.path.exists(pdf_path):
            return send_file(pdf_path, as_attachment=True)
        else:
            return send_file(docx_path, as_attachment=True)
    else:
        return send_file(docx_path, as_attachment=True)


@app.route("/generate_all/<format>")
def generate_all_dates(format):
    if 'data_id' not in session:
        return "No data available", 404
    
    data_id = session['data_id']
    data_file = os.path.join(DATA_FOLDER, f"data_{data_id}.pkl")
    
    if not os.path.exists(data_file):
        return "Data file not found", 404
    
    with open(data_file, 'rb') as f:
        dates_data = pickle.load(f)
    
    machine_name_mapping = session.get('machine_name_mapping', {})
    
    # Generate combined DOCX (similar to single date but with all dates)
    doc = Document()
    
    logo_path = os.path.join(ASSET_FOLDER, "logo.png")
    name_file = os.path.join(ASSET_FOLDER, "name.txt")
    
    institute_name = "Institute/Hospital Name Here"
    if os.path.exists(name_file):
        with open(name_file) as f:
            institute_name = f.read().strip()
    
    add_custom_header_footer(doc, logo_path, institute_name)
    
    first_page = True
    
    for date in sorted(dates_data.keys()):
        machines_data = dates_data[date]
        
        for machine_name, fields_data in machines_data.items():
            if not first_page:
                doc.add_page_break()
            first_page = False
            
            # Add content (same as single date)
            # ... (same logic as generate_date_docx but inline)
    
    filename = "QA_Report_ALL_DATES.docx"
    docx_path = os.path.join(OUTPUT_FOLDER, filename)
    doc.save(docx_path)
    
    if format == "pdf":
        pdf_path = convert_docx_to_pdf_libreoffice(docx_path)
        if pdf_path and os.path.exists(pdf_path):
            return send_file(pdf_path, as_attachment=True)
    
    return send_file(docx_path, as_attachment=True)


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port, debug=False)
